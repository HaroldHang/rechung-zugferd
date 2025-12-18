from pathlib import Path
from docx import Document
from loguru import logger


def extract_docx_text(input_path: Path) -> str:
    logger.info(f"Lese DOCX: {input_path}")
    doc = Document(str(input_path))
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    # Tables
    for table in doc.tables:
        for row in table.rows:
            parts.append("; ".join(cell.text for cell in row.cells))
    return "\n".join(parts)